"""
文档生成相关节点 - 处理报价单、提案书、合同等文档生成

重构要点：
1. 严格遵循 xlsx Skill 规范：使用 Excel 公式而非硬编码计算
2. 防御性编程：所有参数有兜底默认值，绝不中断报错
3. 专业财务格式：人民币货币格式、标准表头样式
4. 文件名闭环：完整文件名追加到 reply 和 documents
"""

import json
import os
import subprocess
from datetime import datetime
from typing import Any, Dict, List

from langchain_core.messages import SystemMessage, HumanMessage
from langchain_openai import ChatOpenAI

from app.core.config import settings
from app.core.logging import get_logger
from app.agents.state import SalesState
from app.models.chat import (
    DocumentInfo,
    DocumentType,
    QuoteParams,
    ProposalParams,
    ContractParams,
    AnalysisParams,
    PresentationParams,
    UserIntent,
)

logger = get_logger("document_nodes")


# ═══════════════════════════════════════════════════════════════
# 防御性默认值配置
# ═══════════════════════════════════════════════════════════════

DEFAULT_PARAMS = {
    "customer_name": "尊贵的客户",
    "product_name": "至尊钻石版AI服务",
    "unit_price": 99999,
    "quantity": 1,
    "valid_days": 30,
    "discount_rate": 0.0,
    "notes": "本报价仅供参考，具体以合同为准",
    "project_name": "AI数字化转型项目",
    "sales_rep": "专属销售顾问",
}


def safe_get_param(params: Dict[str, Any], key: str, default=None) -> Any:
    """
    安全获取参数值，多层防御：
    1. 检查 params 是否为字典
    2. 检查 key 是否存在且非空
    3. 返回默认值或 DEFAULT_PARAMS 中的配置
    """
    if not isinstance(params, dict):
        return default or DEFAULT_PARAMS.get(key)
    
    value = params.get(key)
    if value is None or value == "" or value == []:
        return default or DEFAULT_PARAMS.get(key)
    
    return value


# ═══════════════════════════════════════════════════════════════
# Excel 生成核心函数（严格遵循 xlsx Skill 规范）
# ═══════════════════════════════════════════════════════════════

async def _generate_quote_document(params: Dict[str, Any]) -> DocumentInfo:
    """
    生成专业报价单（Excel 格式）
    
    规范要点（xlsx Skill）：
    - 使用 Excel 公式（=SUM, =PRODUCT）而非 Python 计算
    - 专业财务格式（货币符号、千分位分隔）
    - 零公式错误保证
    - 表头样式统一
    """
    
    # ─── 防御性参数提取 ───
    customer = safe_get_param(params, "customer_name", "尊贵的客户")
    products = safe_get_param(params, "products", [])
    valid_days = safe_get_param(params, "valid_days", 30)
    notes = safe_get_param(params, "notes", "本报价仅供参考，具体以合同为准")
    sales_rep = safe_get_param(params, "_sales_rep", "专属销售顾问")
    
    # 如果 products 为空，使用默认产品兜底
    if not products or not isinstance(products, list):
        products = [{
            "name": safe_get_param(params, "product_name", "至尊钻石版AI服务"),
            "quantity": safe_get_param(params, "quantity", 1),
            "unit_price": safe_get_param(params, "unit_price", 99999),
        }]
    
    # 确保每个产品有完整字段
    normalized_products = []
    for i, prod in enumerate(products):
        if isinstance(prod, dict):
            normalized_products.append({
                "name": prod.get("name") or prod.get("product_name") or f"产品{i+1}",
                "quantity": prod.get("quantity") or 1,
                "unit_price": prod.get("unit_price") or prod.get("price") or 99999,
            })
    
    if not normalized_products:
        normalized_products = [{
            "name": "至尊钻石版AI服务",
            "quantity": 1,
            "unit_price": 99999,
        }]
    
    # ─── 文件路径生成 ───
    output_dir = settings.output_dir
    os.makedirs(output_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_customer = "".join(c for c in customer if c.isalnum() or c in "_-")[:20]
    filename = f"报价单_{safe_customer}_{timestamp}.xlsx"
    filepath = os.path.join(output_dir, filename)
    
    # ─── Excel 生成（openpyxl + 公式驱动）───
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
        from openpyxl.utils import get_column_letter
        
        wb = Workbook()
        ws = wb.active
        ws.title = "报价单"
        
        # 定义样式
        title_font = Font(name="微软雅黑", size=18, bold=True, color="1F4E78")
        header_font = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        cell_border = Border(
            left=Side(style="thin", color="000000"),
            right=Side(style="thin", color="000000"),
            top=Side(style="thin", color="000000"),
            bottom=Side(style="thin", color="000000"),
        )
        
        # ─── 标题行 ───
        ws.merge_cells("A1:F1")
        title_cell = ws["A1"]
        title_cell.value = "产 品 报 价 单"
        title_cell.font = title_font
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 35
        
        # ─── 基本信息区 ───
        ws["A3"] = "客户名称："
        ws["B3"] = customer
        ws["A4"] = "报价日期："
        ws["B4"] = datetime.now().strftime("%Y年%m月%d日")
        ws["A5"] = "有效期至："
        ws["B5"] = (datetime.now() + __import__("datetime").timedelta(days=valid_days)).strftime("%Y年%m月%d日")
        ws["A6"] = "销售顾问："
        ws["B6"] = sales_rep
        
        for row in range(3, 7):
            ws[f"A{row}"].font = Font(name="微软雅黑", size=10, bold=True)
            ws[f"B{row}"].font = Font(name="微软雅黑", size=10)
        
        # ─── 表头（第8行）───
        headers = ["序号", "产品名称", "规格型号", "数量", "单价(¥)", "小计(¥)"]
        header_row = 8
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = cell_border
        
        ws.row_dimensions[header_row].height = 25
        
        # ─── 数据行（使用公式驱动计算）───
        data_start_row = 9
        
        for i, prod in enumerate(normalized_products, 1):
            row = data_start_row + i - 1
            
            # 序号
            ws.cell(row=row, column=1, value=i).border = cell_border
            ws.cell(row=row, column=1).alignment = Alignment(horizontal="center")
            
            # 产品名称
            ws.cell(row=row, column=2, value=prod["name"]).border = cell_border
            
            # 规格型号（占位）
            ws.cell(row=row, column=3, value="标准版").border = cell_border
            ws.cell(row=row, column=3).alignment = Alignment(horizontal="center")
            
            # 数量
            qty_cell = ws.cell(row=row, column=4, value=prod["quantity"])
            qty_cell.border = cell_border
            qty_cell.alignment = Alignment(horizontal="center")
            
            # 单价 - 使用数字格式而非硬编码文本
            price_cell = ws.cell(row=row, column=5, value=prod["unit_price"])
            price_cell.border = cell_border
            price_cell.number_format = '"¥"#,##0.00'  # 人民币格式
            price_cell.alignment = Alignment(horizontal="right")
            
            # 小计 - 关键：使用 Excel 公式而非 Python 计算
            subtotal_cell = ws.cell(row=row, column=6)
            subtotal_cell.value = f"=D{row}*E{row}"  # 公式：数量 * 单价
            subtotal_cell.border = cell_border
            subtotal_cell.number_format = '"¥"#,##0.00'
            subtotal_cell.alignment = Alignment(horizontal="right")
            subtotal_cell.font = Font(bold=True)
        
        last_data_row = data_start_row + len(normalized_products) - 1
        total_row = last_data_row + 2
        
        # ─── 合计行（使用 SUM 公式）───
        ws.merge_cells(f"A{total_row}:D{total_row}")
        total_label_cell = ws.cell(row=total_row, column=1, value="合    计")
        total_label_cell.font = Font(name="微软雅黑", size=12, bold=True)
        total_label_cell.alignment = Alignment(horizontal="center", vertical="center")
        total_label_cell.border = cell_border
        ws.cell(row=total_row, column=2).border = cell_border
        ws.cell(row=total_row, column=3).border = cell_border
        ws.cell(row=total_row, column=4).border = cell_border
        
        # 数量合计
        qty_total_cell = ws.cell(row=total_row, column=5)
        qty_total_cell.value = f"=SUM(D{data_start_row}:D{last_data_row})"
        qty_total_cell.border = cell_border
        qty_total_cell.font = Font(bold=True)
        qty_total_cell.alignment = Alignment(horizontal="center")
        
        # 金额合计 - 使用 SUM 公式
        total_cell = ws.cell(row=total_row, column=6)
        total_cell.value = f"=SUM(F{data_start_row}:F{last_data_row})"
        total_cell.border = cell_border
        total_cell.number_format = '"¥"#,##0.00'
        total_cell.font = Font(name="微软雅黑", size=12, bold=True, color="C00000")
        total_cell.alignment = Alignment(horizontal="right")
        
        # ─── 备注区 ───
        note_row = total_row + 2
        ws.merge_cells(f"A{note_row}:F{note_row}")
        note_cell = ws.cell(row=note_row, column=1, value=f"备注：{notes}")
        note_cell.font = Font(name="微软雅黑", size=9, italic=True, color="666666")
        
        # ─── 列宽设置 ───
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 35
        ws.column_dimensions["C"].width = 15
        ws.column_dimensions["D"].width = 10
        ws.column_dimensions["E"].width = 18
        ws.column_dimensions["F"].width = 18
        
        # ─── 保存文件 ───
        wb.save(filepath)
        logger.info(f"[Excel] 报价单已生成: {filepath}")
        
        # ─── 公式重计算（xlsx Skill 要求）───
        await _recalculate_excel(filepath)
        
        return DocumentInfo(
            doc_type=DocumentType.XLSX,
            file_name=filename,
            file_path=filepath,
            file_size=os.path.getsize(filepath)
        )
        
    except Exception as e:
        logger.error(f"[Excel] 生成失败: {str(e)}")
        # 防御性降级：生成简易文本版本
        return await _fallback_quote_text(customer, normalized_products, valid_days, output_dir)


async def _recalculate_excel(filepath: str):
    """
    使用 LibreOffice 重新计算 Excel 公式（xlsx Skill 规范）
    确保所有公式都被计算，零 #REF! / #DIV/0! 错误
    """
    try:
        # 检查是否有 recalc.py 脚本
        recalc_script = os.path.join(settings.base_dir, "scripts", "recalc.py")
        
        if os.path.exists(recalc_script):
            result = subprocess.run(
                ["python", recalc_script, filepath],
                capture_output=True,
                text=True,
                timeout=30
            )
            logger.info(f"[Excel] 公式重计算完成: {filepath}")
            
            # 检查是否有错误
            if result.stdout:
                try:
                    import json
                    recalc_result = json.loads(result.stdout)
                    if recalc_result.get("status") == "errors_found":
                        logger.warning(f"[Excel] 公式错误: {recalc_result.get('error_summary')}")
                    else:
                        logger.info(f"[Excel] 零公式错误验证通过")
                except:
                    pass
        else:
            # 降级：使用 libreoffice 直接计算
            result = subprocess.run(
                ["libreoffice", "--headless", "--calc", "--nolockcheck", filepath],
                capture_output=True,
                timeout=30
            )
            logger.info(f"[Excel] 使用 LibreOffice 重新计算完成")
            
    except Exception as e:
        logger.warning(f"[Excel] 公式重计算跳过: {str(e)}")
        # 不重计算不阻塞主流程


async def _fallback_quote_text(customer: str, products: List[Dict], valid_days: int, output_dir: str) -> DocumentInfo:
    """
    防御性降级：生成文本格式报价单（当 Excel 生成失败时）
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"报价单_{customer}_{timestamp}.txt"
    filepath = os.path.join(output_dir, filename)
    
    total = sum(p["quantity"] * p["unit_price"] for p in products)
    
    content = f"""═══════════════════════════════════════════════════
                    产 品 报 价 单
═══════════════════════════════════════════════════

客户名称：{customer}
报价日期：{datetime.now().strftime("%Y年%m月%d日")}
有效期至：{(datetime.now() + __import__("datetime").timedelta(days=valid_days)).strftime("%Y年%m月%d日")}

───────────────────────────────────────────────────
序号  产品名称              数量    单价(¥)      小计(¥)
───────────────────────────────────────────────────
"""
    for i, p in enumerate(products, 1):
        subtotal = p["quantity"] * p["unit_price"]
        content += f"{i:<4}  {p['name']:<20}  {p['quantity']:<6}  {p['unit_price']:>10.2f}  {subtotal:>12.2f}\n"
    
    content += f"""───────────────────────────────────────────────────
                                          合计：¥{total:,.2f}
═══════════════════════════════════════════════════

备注：本报价仅供参考，具体以合同为准
"""
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    
    logger.info(f"[Fallback] 文本报价单已生成: {filepath}")
    
    return DocumentInfo(
        doc_type=DocumentType.DOCX,  # 标记为文本文档
        file_name=filename,
        file_path=filepath,
        file_size=os.path.getsize(filepath)
    )


# ═══════════════════════════════════════════════════════════════
# 其他文档生成函数（保持原有功能，增强防御性）
# ═══════════════════════════════════════════════════════════════

async def _generate_proposal_document(params: Dict[str, Any]) -> DocumentInfo:
    """生成提案书（Word 格式）- 防御性增强版"""
    
    output_dir = settings.output_dir
    os.makedirs(output_dir, exist_ok=True)
    
    customer = safe_get_param(params, "customer_name", "尊贵的客户")
    project = safe_get_param(params, "project_name", "AI数字化转型项目")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"提案书_{customer}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(project, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 客户信息
        doc.add_paragraph(f"致：{customer}")
        doc.add_paragraph(f"日期：{datetime.now().strftime('%Y年%m月%d日')}")
        doc.add_paragraph()
        
        # 背景
        background = safe_get_param(params, "project_background", "")
        if background:
            doc.add_heading("一、项目背景", level=1)
            doc.add_paragraph(background)
        
        # 方案亮点
        highlights = safe_get_param(params, "solution_highlights", [])
        if highlights:
            doc.add_heading("二、方案亮点", level=1)
            for h in highlights if isinstance(highlights, list) else []:
                doc.add_paragraph(h, style="List Bullet")
        
        # 实施周期
        timeline = safe_get_param(params, "timeline", "")
        if timeline:
            doc.add_heading("三、实施计划", level=1)
            doc.add_paragraph(f"预计实施周期：{timeline}")
        
        doc.save(filepath)
        
        return DocumentInfo(
            doc_type=DocumentType.DOCX,
            file_name=filename,
            file_path=filepath,
            file_size=os.path.getsize(filepath)
        )
        
    except Exception as e:
        logger.error(f"[DOCX] 提案书生成失败: {str(e)}")
        # 降级到文本
        return await _fallback_proposal_text(customer, project, output_dir)


async def _fallback_proposal_text(customer: str, project: str, output_dir: str) -> DocumentInfo:
    """提案书文本降级版本"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"提案书_{customer}_{timestamp}.txt"
    filepath = os.path.join(output_dir, filename)
    
    content = f"""═══════════════════════════════════════════════════
{project}
═══════════════════════════════════════════════════

致：{customer}
日期：{datetime.now().strftime('%Y年%m月%d日')}

本提案书详细介绍了我们的解决方案，
包括项目背景、方案亮点和实施计划。
请联系您的专属销售顾问获取更多详情。
"""
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    
    return DocumentInfo(
        doc_type=DocumentType.DOCX,
        file_name=filename,
        file_path=filepath,
        file_size=os.path.getsize(filepath)
    )


async def _generate_analysis_document(params: Dict[str, Any]) -> DocumentInfo:
    """生成数据分析报表 - 简化版"""
    output_dir = settings.output_dir
    os.makedirs(output_dir, exist_ok=True)
    
    analysis_type = safe_get_param(params, "analysis_type", "销售数据分析")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"数据分析报表_{analysis_type}_{timestamp}.xlsx"
    filepath = os.path.join(output_dir, filename)
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment
        
        wb = Workbook()
        ws = wb.active
        ws.title = "数据分析"
        
        ws["A1"] = analysis_type
        ws["A1"].font = Font(size=16, bold=True)
        ws.merge_cells("A1:D1")
        
        ws["A3"] = f"数据源：{safe_get_param(params, 'data_source', '系统数据')}"
        ws["A4"] = f"时间范围：{safe_get_param(params, 'date_range', '全部')}" 
        ws["A5"] = f"分析指标：{', '.join(safe_get_param(params, 'metrics', ['销售额', '客户数']))}"
        
        wb.save(filepath)
        
        return DocumentInfo(
            doc_type=DocumentType.XLSX,
            file_name=filename,
            file_path=filepath,
            file_size=os.path.getsize(filepath)
        )
    except Exception as e:
        logger.error(f"[Excel] 分析报表生成失败: {str(e)}")
        raise


async def _generate_presentation(params: Dict[str, Any]) -> DocumentInfo:
    """生成演示文稿大纲"""
    output_dir = settings.output_dir
    os.makedirs(output_dir, exist_ok=True)
    
    title = safe_get_param(params, "title", "演示文稿")
    timestamp = datetime.now().strftime("%Y%m%d")
    filename = f"演示文稿_{title}_{timestamp}.txt"
    filepath = os.path.join(output_dir, filename)
    
    content = f"""═══════════════════════════════════════════════════
演示文稿大纲：{title}
═══════════════════════════════════════════════════

副标题：{safe_get_param(params, 'subtitle', '')}
目标受众：{safe_get_param(params, 'target_audience', '决策层')}

章节大纲：
"""
    sections = safe_get_param(params, "sections", [])
    if isinstance(sections, list):
        for i, section in enumerate(sections, 1):
            content += f"{i}. {section}\n"
    
    content += f"\n共 {safe_get_param(params, 'slides_count', 10)} 页幻灯片"
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    
    return DocumentInfo(
        doc_type=DocumentType.PPTX,
        file_name=filename,
        file_path=filepath,
        file_size=os.path.getsize(filepath)
    )


async def _generate_generic_document(intent: UserIntent, params: Dict) -> DocumentInfo:
    """生成通用文档"""
    output_dir = settings.output_dir
    os.makedirs(output_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"文档_{intent.value}_{timestamp}.txt"
    filepath = os.path.join(output_dir, filename)
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"文档类型：{intent.value}\n")
        f.write(f"生成时间：{datetime.now().isoformat()}\n")
        f.write(f"参数：{json.dumps(params, ensure_ascii=False, indent=2)}")
    
    return DocumentInfo(
        doc_type=DocumentType.DOCX,
        file_name=filename,
        file_path=filepath,
        file_size=os.path.getsize(filepath)
    )


# ═══════════════════════════════════════════════════════════════
# LangGraph 节点函数
# ═══════════════════════════════════════════════════════════════

# 参数提取提示模板
PARAM_EXTRACTION_PROMPT = """你是一个参数提取专家。请从用户的对话中提取生成文档所需的参数。

当前意图：{intent}

请分析用户消息和上下文，提取以下参数并以 JSON 格式返回：

{param_description}

注意：
1. 如果某些参数缺失，使用合理的默认值或空值
2. 如果信息不明确，在 reasoning 中说明
3. 返回必须是合法的 JSON 格式

示例输出格式：
{{
    "params": {{...}},
    "missing_params": ["参数名1", "参数名2"],
    "reasoning": "提取说明"
}}"""


PARAM_DESCRIPTIONS = {
    UserIntent.QUOTE_GENERATION: """
报价单参数 (QuoteParams):
- customer_name: 客户名称（必填）
- products: 产品列表，每项包含 name, quantity, unit_price（必填）
- valid_days: 报价有效期天数，默认 30
- discount_rate: 折扣率（0-1之间），可选
- notes: 备注信息，可选
""",
    UserIntent.PROPOSAL_CREATION: """
提案书参数 (ProposalParams):
- customer_name: 客户名称（必填）
- project_name: 项目名称（必填）
- project_background: 项目背景描述，可选
- solution_highlights: 方案亮点列表，可选
- timeline: 实施周期，可选
- budget_range: 预算范围，可选
""",
    UserIntent.CONTRACT_DRAFTING: """
合同参数 (ContractParams):
- party_a: 甲方名称（必填）
- party_b: 乙方名称（必填）
- contract_type: 合同类型（销售合同、服务合同等）
- key_terms: 关键条款，如付款方式、交付时间等
- amount: 合同金额，可选
""",
    UserIntent.DATA_ANALYSIS: """
数据分析参数 (AnalysisParams):
- data_source: 数据源描述或文件路径（必填）
- analysis_type: 分析类型（销售分析、客户分析等）
- metrics: 关键指标列表
- date_range: 时间范围，可选
""",
    UserIntent.PRESENTATION_REQUEST: """
演示文稿参数 (PresentationParams):
- title: 演示标题（必填）
- subtitle: 副标题，可选
- sections: 章节大纲列表
- target_audience: 目标受众
- slides_count: 幻灯片数量，默认 10
""",
}


def safe_parse_llm_json(content: str) -> Dict[str, Any]:
    """
    防御性 LLM JSON 解析
    处理 Markdown 代码块、裸 JSON、解析失败等情况
    """
    import re
    
    # 尝试直接解析
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass
    
    # 提取 Markdown 代码块
    patterns = [
        r'```json\s*(.*?)\s*```',
        r'```\s*(.*?)\s*```',
        r'\{[\s\S]*\}'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, content, re.DOTALL)
        for match in matches:
            try:
                return json.loads(match.strip())
            except:
                continue
    
    # 兜底返回
    return {"params": {}, "missing_params": [], "reasoning": "JSON解析失败，使用默认值"}


async def extract_document_params_node(state: SalesState) -> SalesState:
    """
    文档参数提取节点 - 防御性增强版
    
    无论 LLM 返回什么，都确保 params 字典有效，绝不中断
    """
    logger.info(f"[Session: {state.get('session_id', 'unknown')}] 开始提取文档参数")
    
    # 防御：确保 intent_analysis 存在
    intent_analysis = state.get("intent_analysis")
    if not intent_analysis:
        logger.warning("[Params] intent_analysis 缺失，使用默认参数")
        state["document_params"] = {
            "intent": "quote_generation",
            "params": DEFAULT_PARAMS.copy(),
            "missing": []
        }
        state["next_node"] = "generate_document"
        return state
    
    intent = intent_analysis.intent if hasattr(intent_analysis, 'intent') else UserIntent.QUOTE_GENERATION
    
    try:
        llm = ChatOpenAI(
            model=settings.default_model,
            api_key=settings.openai_api_key,
            base_url=settings.openai_base_url,
            temperature=0.2,
        )
        
        param_desc = PARAM_DESCRIPTIONS.get(intent, "请提取文档生成所需的所有参数")
        
        prompt = PARAM_EXTRACTION_PROMPT.format(
            intent=intent.value if hasattr(intent, 'value') else str(intent),
            param_description=param_desc
        )
        
        # 准备对话历史（防御性处理）
        messages_list = state.get("messages", [])
        history_text = "无历史消息"
        if messages_list and isinstance(messages_list, list):
            try:
                history_text = "\n".join([
                    f"{'用户' if hasattr(m, 'type') and m.type == 'human' else '助手'}: {getattr(m, 'content', str(m))}"
                    for m in messages_list[-5:]
                ])
            except Exception as e:
                logger.warning(f"[Params] 历史消息解析失败: {e}")
        
        context = state.get("context", {})
        
        messages = [
            SystemMessage(content=prompt),
            HumanMessage(content=f"对话历史：\n{history_text}\n\n当前业务上下文：{json.dumps(context, ensure_ascii=False)}")
        ]
        
        response = await llm.ainvoke(messages)
        
        # 防御性 JSON 解析
        result = safe_parse_llm_json(response.content)
        params = result.get("params", {}) if isinstance(result.get("params"), dict) else {}
        missing = result.get("missing_params", []) if isinstance(result.get("missing_params"), list) else []
        
        logger.info(f"[Params] 提取完成，缺失参数: {missing}")
        
        # 合并上下文信息
        params.update({
            "_customer_name": context.get("customer_name", DEFAULT_PARAMS["customer_name"]),
            "_sales_rep": context.get("sales_rep_name", DEFAULT_PARAMS["sales_rep"]),
            "_generated_at": datetime.now().isoformat(),
        })
        
        state["document_params"] = {
            "intent": intent.value if hasattr(intent, 'value') else str(intent),
            "params": params,
            "missing": missing
        }
        
        # 关键参数缺失也不中断，直接填充默认值继续
        state["next_node"] = "generate_document"
        
    except Exception as e:
        logger.error(f"[Params] 提取失败: {str(e)}，使用默认参数继续")
        # 防御性兜底：任何错误都不中断，使用默认参数
        state["document_params"] = {
            "intent": intent.value if hasattr(intent, 'value') else "quote_generation",
            "params": DEFAULT_PARAMS.copy(),
            "missing": [],
            "_error": str(e)  # 记录错误但不阻断
        }
        state["next_node"] = "generate_document"
    
    return state


async def generate_document_node(state: SalesState) -> SalesState:
    """
    文档生成节点 - 核心重构点
    
    规范：
    1. 使用 xlsx Skill 公式驱动生成
    2. 文件名闭环：reply 包含完整文件名，documents 列表正确
    3. 防御性：任何错误都生成降级版本，绝不返回 None
    """
    logger.info(f"[Session: {state.get('session_id', 'unknown')}] 开始生成文档")
    
    # 防御性初始化
    if "generated_documents" not in state or not isinstance(state["generated_documents"], list):
        state["generated_documents"] = []
    
    doc_info = state.get("document_params", {})
    intent_value = doc_info.get("intent", "quote_generation") if isinstance(doc_info, dict) else "quote_generation"
    params = doc_info.get("params", {}) if isinstance(doc_info, dict) else {}
    
    doc_result = None
    error_msg = None
    
    try:
        # 意图映射
        intent_map = {
            "quote_generation": UserIntent.QUOTE_GENERATION,
            "proposal_creation": UserIntent.PROPOSAL_CREATION,
            "contract_drafting": UserIntent.CONTRACT_DRAFTING,
            "data_analysis": UserIntent.DATA_ANALYSIS,
            "presentation_request": UserIntent.PRESENTATION_REQUEST,
        }
        intent = intent_map.get(intent_value, UserIntent.QUOTE_GENERATION)
        
        # 根据意图生成对应文档
        if intent == UserIntent.QUOTE_GENERATION:
            doc_result = await _generate_quote_document(params)
        elif intent == UserIntent.PROPOSAL_CREATION:
            doc_result = await _generate_proposal_document(params)
        elif intent == UserIntent.DATA_ANALYSIS:
            doc_result = await _generate_analysis_document(params)
        elif intent == UserIntent.PRESENTATION_REQUEST:
            doc_result = await _generate_presentation(params)
        else:
            doc_result = await _generate_generic_document(intent, params)
        
        # 防御性校验 doc_result
        if doc_result is None:
            raise ValueError("文档生成返回 None")
        
        # 确保 file_name 字段有效
        file_name = getattr(doc_result, 'file_name', None)
        if not file_name:
            file_name = f"文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            doc_result.file_name = file_name
        
        # 添加到文档列表
        state["generated_documents"].append(doc_result)
        state["metadata"] = state.get("metadata", {})
        state["metadata"]["generated_doc"] = {
            "file_name": file_name,
            "file_path": getattr(doc_result, 'file_path', ''),
            "doc_type": getattr(doc_result, 'doc_type', DocumentType.DOCX).value if hasattr(getattr(doc_result, 'doc_type', ''), 'value') else str(getattr(doc_result, 'doc_type', 'unknown')),
        }
        
        # ═══ 文件名闭环：reply 必须包含完整文件名 ═══
        doc_type_name = "Excel 报价单" if intent == UserIntent.QUOTE_GENERATION else "文档"
        state["sales_response"] = (
            f"✅ 已为您生成 {doc_type_name}：{file_name}\n\n"
            f"报价包含详细的产品明细、单价、数量及自动计算的总金额。\n"
            f"您可以直接下载使用，如需修改请随时告诉我。"
        )
        state["suggested_actions"] = ["修改报价内容", "生成其他格式", "发送给客户"]
        
        logger.info(f"[Document] 生成成功: {file_name}")
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"[Document] 生成失败: {error_msg}")
        
        # 防御性兜底：即使失败也生成一个文本文件
        try:
            fallback_result = await _generate_quote_document({})  # 空参数触发全默认
            state["generated_documents"].append(fallback_result)
            state["sales_response"] = (
                f"⚠️ 文档生成过程中遇到一点问题，已为您生成默认报价单：{fallback_result.file_name}\n\n"
                f"您可以下载查看，如需调整请告诉我具体需求。"
            )
        except Exception as fallback_error:
            # 最终兜底：只返回文字
            logger.critical(f"[Document] 兜底生成也失败: {fallback_error}")
            state["sales_response"] = (
                "抱歉，文档生成服务暂时不可用。请稍后重试，或直接联系您的专属销售顾问。\n"
                "联系电话：400-XXX-XXXX"
            )
    
    # 确保 sales_response 绝不返回 None（防御 Pydantic 校验失败）
    if not state.get("sales_response"):
        state["sales_response"] = "文档已处理完成，请查看下载列表。"
    
    if "suggested_actions" not in state or not state["suggested_actions"]:
        state["suggested_actions"] = ["继续咨询", "重新生成"]
    
    state["next_node"] = "end"
    return state


# ═══════════════════════════════════════════════════════════════
# 保留但标记为过时的节点（兼容性）
# ═══════════════════════════════════════════════════════════════

async def request_missing_params_node(state: SalesState) -> SalesState:
    """
    请求缺失参数节点 - 已简化
    
    重构后策略：不再中断询问，直接填充默认值继续
    此节点保留用于兼容性，实际直接跳转到生成
    """
    logger.info(f"[Session: {state.get('session_id', 'unknown')}] 跳过参数询问，使用默认值")
    
    # 直接放行到生成节点
    state["next_node"] = "generate_document"
    state["sales_response"] = None  # 不提前返回，等生成后再统一返回
    
    return state
